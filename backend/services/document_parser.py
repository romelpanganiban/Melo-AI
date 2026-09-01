"""Extract text from supported document file formats."""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

from core.errors import ChatServiceError, ValidationError
from core.settings import settings


def sanitize_filename(filename: str) -> str:
    """Normalize a user-provided filename into a safe, database-friendly value."""
    if filename is None:
        raise ValidationError("filename is required")

    raw_name = str(filename).strip()
    if not raw_name:
        raise ValidationError("filename is required")

    from core.logging import logger
    logger.info(f"[sanitize_filename] Input filename: {repr(filename)}")

    normalized = raw_name.replace("\\", "/")
    logger.info(f"[sanitize_filename] After strip/replace backslash: {repr(normalized)}")

    basename = normalized.split("/")[-1]
    logger.info(f"[sanitize_filename] After split by /: {repr(basename)}")

    basename = basename.replace("\x00", "")
    basename = re.sub(r"[<>:\"|?*\r\n\t]", "_", basename)
    basename = re.sub(r"\s+", " ", basename).strip()
    logger.info(f"[sanitize_filename] After regex replacements: {repr(basename)}")

    basename = basename.strip(". ")
    logger.info(f"[sanitize_filename] After strip dots/spaces: {repr(basename)}")

    if not basename or basename in {".", ".."}:
        suffix = Path(raw_name).suffix.lower()
        safe_name = "uploaded-document"
        if suffix:
            safe_name = f"{safe_name}{suffix}"

        logger.warning(f"[sanitize_filename] Falling back to safe name: {repr(safe_name)}")
        return safe_name

    return basename


class DocumentParser:
    """Convert uploaded TXT, PDF, and DOCX files into plain text."""

    SUPPORTED_TYPES = {"txt", "pdf", "docx"}

    def parse(self, filename: str, content: bytes) -> tuple[str, str]:
        safe_name = sanitize_filename(filename)
        if not content:
            raise ValidationError("file is empty")

        file_type = Path(safe_name).suffix.lower().lstrip(".")
        if file_type not in self.SUPPORTED_TYPES:
            raise ValidationError("Only .txt, .pdf, and .docx files are supported")

        if file_type == "txt":
            text = self._decode_text_file(content)
        elif file_type == "pdf":
            text = self._parse_pdf(content)
        else:
            text = self._parse_docx(content)

        if not text.strip():
            raise ValidationError("The uploaded file contains no extractable text")
        if len(text) > settings.MAX_DOCUMENT_CONTENT_LENGTH:
            raise ValidationError("Extracted document text exceeds the 2 MB limit", field="content")

        return file_type, text

    def _decode_text_file(self, content: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-16", "utf-16-le", "utf-16-be", "cp1252", "latin-1"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ValidationError("The text file could not be decoded. Please save it as UTF-8 or UTF-16 text and try again.")

    def _parse_pdf(self, content: bytes) -> str:
        try:
            from pypdf import PdfReader

            reader = PdfReader(BytesIO(content))
            parts = []
            total_length = 0
            for page in reader.pages:
                page_text = page.extract_text() or ""
                total_length += len(page_text)
                if total_length > settings.MAX_DOCUMENT_CONTENT_LENGTH:
                    raise ValidationError("Extracted PDF text exceeds the 2 MB limit", field="content")
                parts.append(page_text)
            return "\n".join(parts)
        except Exception as exc:
            raise ChatServiceError(f"Failed to extract PDF text: {exc}") from exc

    def _parse_docx(self, content: bytes) -> str:
        try:
            from docx import Document

            document = Document(BytesIO(content))
            paragraphs = []
            total_length = 0
            for paragraph in document.paragraphs:
                total_length += len(paragraph.text)
                if total_length > settings.MAX_DOCUMENT_CONTENT_LENGTH:
                    raise ValidationError("Extracted DOCX text exceeds the 2 MB limit", field="content")
                paragraphs.append(paragraph.text)
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    total_length += len(row_text)
                    if total_length > settings.MAX_DOCUMENT_CONTENT_LENGTH:
                        raise ValidationError("Extracted DOCX text exceeds the 2 MB limit", field="content")
                    paragraphs.append(row_text)
            return "\n".join(paragraphs)
        except Exception as exc:
            raise ChatServiceError(f"Failed to extract DOCX text: {exc}") from exc


_document_parser = DocumentParser()


def get_document_parser() -> DocumentParser:
    return _document_parser
